"""One-click apply: build applicant answers, render the CV, drive the ATS submit."""
import io

from ada.db.models import ApplicationStatus, Job, Profile, RunStatus, User
from ada.db.repositories import (
    ApplicationRepository,
    JobRepository,
    ProfileRepository,
    RunRepository,
)
from ada.db.session import _session_factory
from ada.observability import log
from ada.services.ats import ApplicantAnswers, FormPlan, split_name
from ada.services.ats import ashby as ats_ashby
from ada.services.ats import greenhouse as ats_greenhouse
from ada.services.ats import lever as ats_lever

SUPPORTED_SOURCES = {"greenhouse", "lever", "ashby"}


class ApplyPrecondition(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


def cv_docx_bytes(cv_markdown: str) -> bytes:
    from docx import Document

    doc = Document()
    for raw in cv_markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.lstrip("#").strip()
        if line.startswith("# "):
            doc.add_heading(stripped, level=1)
        elif line.startswith("## "):
            doc.add_heading(stripped, level=2)
        elif line.startswith("### "):
            doc.add_heading(stripped, level=3)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line.replace("**", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_answers(user: User, profile: Profile | None, cv_markdown: str) -> ApplicantAnswers:
    if profile is None or not (profile.full_name or "").strip():
        raise ApplyPrecondition(
            "identity_required", "Add your full name (and phone) so Ada can apply for you."
        )
    full_name = (profile.full_name or "").strip()
    first, last = split_name(full_name)
    return ApplicantAnswers(
        full_name=full_name,
        first_name=first,
        last_name=last,
        email=user.email,
        phone=(profile.phone or "").strip() or None,
        linkedin_url=profile.linkedin_url,
        cv_filename=f"{full_name.replace(' ', '-')}-CV.docx",
        cv_bytes=cv_docx_bytes(cv_markdown),
    )


def build_plan(job: Job, answers: ApplicantAnswers) -> FormPlan:
    if not job.url:
        raise ApplyPrecondition("no_url", "This listing has no application page on file.")
    if job.source == ats_greenhouse.SOURCE:
        return ats_greenhouse.build_form_plan(job.url, answers)
    if job.source == ats_lever.SOURCE:
        return ats_lever.build_form_plan(job.url, answers)
    if job.source == ats_ashby.SOURCE:
        return ats_ashby.build_form_plan(job.url, answers)
    raise ApplyPrecondition("unsupported_source", "No deterministic plan for this source.")


async def latest_cv_markdown(runs: RunRepository, user_id: str) -> tuple[str, str] | None:
    for run in await runs.list_by_user(user_id):
        if run.status == RunStatus.COMPLETE and run.rewritten_cv:
            return run.rewritten_cv, run.id
    return None


async def run_submission(application_id: str) -> None:
    from ada.services.ats.executor import execute_plan
    from ada.services.ats.generic import submit_generic

    async with _session_factory() as session:
        applications = ApplicationRepository(session)
        application = await applications.get(application_id)
        if application is None:
            return
        job = await JobRepository(session).get(application.job_id)
        profile = await ProfileRepository(session).get(application.user_id)
        user = await session.get(User, application.user_id)
        run = await RunRepository(session).get(application.run_id) if application.run_id else None
        try:
            if user is None or job is None or run is None or not run.rewritten_cv:
                raise ApplyPrecondition("missing_data", "Application inputs went missing.")
            if not job.url:
                raise ApplyPrecondition("no_url", "This listing has no application page on file.")
            answers = build_answers(user, profile, run.rewritten_cv)
            plan = build_plan(job, answers) if job.source in SUPPORTED_SOURCES else None
            job_url = job.url
        except ApplyPrecondition as exc:
            await applications.set_status(
                application_id, ApplicationStatus.NEEDS_ATTENTION, detail=str(exc)
            )
            return
    try:
        if plan is not None:
            outcome = await execute_plan(plan, answers)
        else:
            outcome = await submit_generic(job_url, answers)
    except Exception as exc:  # noqa: BLE001 — any submit failure must land in status, not logs alone
        log.error("apply_submit_crashed", application_id=application_id, error=str(exc))
        outcome = None
    async with _session_factory() as session:
        applications = ApplicationRepository(session)
        if outcome is None:
            await applications.set_status(
                application_id, ApplicationStatus.FAILED,
                detail="Something broke on our side — we'll look into it. Nothing was submitted.",
            )
        elif outcome.status == "submitted":
            await applications.set_status(application_id, ApplicationStatus.SUBMITTED)
            log.info("apply_submitted", application_id=application_id)
        else:
            await applications.set_status(
                application_id, ApplicationStatus.NEEDS_ATTENTION, detail=outcome.detail
            )


def is_supported(job: Job) -> bool:
    return bool(job.url)
